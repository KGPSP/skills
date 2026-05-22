#!/usr/bin/env python3
"""Render a PZP procedural letter from a filled `.md` file into a DOCX based on
the KG PSP EZD template `wzor_pismo_przewodnie.docx`.

Usage:
    python3 render_docx.py \
        --template /path/to/wzor_pismo_przewodnie.docx \
        --input    /path/to/W03-wyjasnienia-tresci-oferty-<slug>.md \
        --output   /path/to/W03-wyjasnienia-tresci-oferty-<slug>.docx

The `.md` input must contain:
    1. YAML frontmatter with mandatory fields (see templates/_frontmatter-base.md).
    2. A "## Treść pisma" H2 section containing the letter body as markdown.
    3. "## Załączniki" and "## Otrzymują" tail sections (used to populate the
       corresponding sections in the DOCX).

The script:
    * Copies the DOCX template to the output path (template is never modified).
    * Replaces text within EZD bookmarks:
        ezdSprawaZnak, ezdDataPodpisu, ezdPracownikNazwa,
        ezdPracownikAtrybut1 (stanowisko), ezdPracownikAtrybut2 (tytuł),
        ezdPracownikAtrybut3 (imię i nazwisko — fallback if unused in template).
        ezdAutorWydzialOpis is left intact (constant "Biuro Informatyki i
        Łączności" from the template).
    * Replaces `$`-placeholders in body paragraphs:
        "$sygnatura pisma", "$DataPodpisu", "$stanowisko", "$tytuł",
        "$imię i nazwisko".
    * Replaces the `[adresat]/[jednostka organizacyjna PSP z listy adresatów]`
      paragraph with the real wykonawca name + address.
    * Replaces the two body placeholder paragraphs (Wstęp… / Rozwinięcie…) with
      the actual letter body parsed from the markdown between "## Treść pisma"
      and the tail.
    * Replaces the załączniki placeholder paragraph and adds Otrzymują entries.

Dependencies: python-docx. PyYAML is used if available, otherwise a minimal
YAML parser handles the subset of frontmatter used by the skill.
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as exc:
    sys.stderr.write(
        "ERROR: python-docx is required. Install with `pip install python-docx`.\n"
    )
    raise SystemExit(1) from exc

try:
    import yaml  # type: ignore
    HAS_YAML = True
except ImportError:
    HAS_YAML = False


POLISH_MONTHS = {
    1: "stycznia", 2: "lutego", 3: "marca", 4: "kwietnia",
    5: "maja", 6: "czerwca", 7: "lipca", 8: "sierpnia",
    9: "września", 10: "października", 11: "listopada", 12: "grudnia",
}

BODY_PLACEHOLDER_PREFIXES = (
    "Wstęp do pisma",
    "Rozwinięcie i zasadnicza treść pisma",
)

ATTACHMENT_PLACEHOLDER_PREFIXES = (
    "Opis załącznika",
)

RECIPIENT_PLACEHOLDER_PREFIXES = (
    "Adresat (tytuł, imię, nazwisko, stanowisko, adres)",
)


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def parse_md(path: Path) -> tuple[dict, dict[str, str]]:
    """Return (frontmatter_dict, sections_dict).

    sections keys: "tresc" (everything between "## Treść pisma" and "## Załączniki"
    minus the first metadata table), "zalaczniki", "otrzymuja".
    """
    text = path.read_text(encoding="utf-8")
    fm, body = _split_frontmatter(text)
    sections = _split_sections(body)
    return fm, sections


def _split_frontmatter(text: str) -> tuple[dict, str]:
    if not text.startswith("---"):
        raise ValueError("Missing YAML frontmatter (expected leading ---).")
    end = text.find("\n---", 3)
    if end == -1:
        raise ValueError("Unterminated YAML frontmatter.")
    fm_text = text[3:end].strip("\n")
    body = text[end + 4 :].lstrip("\n")
    fm = _parse_yaml(fm_text)
    return fm, body


def _parse_yaml(text: str) -> dict:
    if HAS_YAML:
        result = yaml.safe_load(text) or {}
        if not isinstance(result, dict):
            raise ValueError("Frontmatter root must be a mapping.")
        return result
    # Minimal fallback parser for the restricted frontmatter subset used by the
    # skill (flat keys, lists as indented dashes, no nested mappings).
    #
    # Splits each top-level line on the first colon that follows the YAML key
    # (matched via regex to avoid partitioning on colons that appear inside
    # quoted values — e.g. `postepowanie: "B10: HPC/AI dla SOiA"`).
    key_line_re = re.compile(r"^([A-Za-z_][A-Za-z0-9_\-]*)\s*:\s*(.*)$")
    out: dict = {}
    current_list: list | None = None
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line or line.lstrip().startswith("#"):
            continue
        indent = len(line) - len(line.lstrip(" "))
        stripped = line.lstrip(" ")
        if indent == 0:
            match = key_line_re.match(stripped)
            if match:
                key = match.group(1)
                value = match.group(2).strip()
                # Strip trailing inline comment (YAML: ' #...').
                hash_pos = _find_inline_comment(value)
                if hash_pos >= 0:
                    value = value[:hash_pos].rstrip()
                if value == "":
                    current_list = []
                    out[key] = current_list
                else:
                    current_list = None
                    out[key] = _coerce_scalar(value)
                continue
        if stripped.startswith("- ") and current_list is not None:
            item = stripped[2:].strip()
            hash_pos = _find_inline_comment(item)
            if hash_pos >= 0:
                item = item[:hash_pos].rstrip()
            current_list.append(_coerce_scalar(item))
    return out


def _find_inline_comment(value: str) -> int:
    """Return index of a YAML inline comment (` #`) or -1 if none.

    Skips `#` that appear inside single- or double-quoted strings so that
    values such as `"B10: HPC/AI"` or `example #1` stay intact.
    """
    in_single = False
    in_double = False
    prev = ""
    for i, ch in enumerate(value):
        if ch == "'" and not in_double:
            in_single = not in_single
        elif ch == '"' and not in_single:
            in_double = not in_double
        elif ch == "#" and not in_single and not in_double and prev in (" ", "\t"):
            return i - 1 if prev in (" ", "\t") else i
        prev = ch
    return -1


def _coerce_scalar(value: str):
    value = value.strip()
    if value.startswith('"') and value.endswith('"'):
        return value[1:-1]
    if value.startswith("'") and value.endswith("'"):
        return value[1:-1]
    if value.lower() in {"true", "false"}:
        return value.lower() == "true"
    if re.fullmatch(r"-?\d+", value):
        return int(value)
    if re.fullmatch(r"-?\d+\.\d+", value):
        return float(value)
    return value


def _split_sections(body: str) -> dict[str, str]:
    # Drop everything before "## Treść pisma" (including data-do-szablonu table).
    tresc_marker = "## Treść pisma"
    if tresc_marker not in body:
        raise ValueError(
            f"Missing required section heading: {tresc_marker!r}. "
            "Letter markdown must include '## Treść pisma' section."
        )
    body = body[body.index(tresc_marker) + len(tresc_marker):]
    # Split by the remaining H2 sections.
    parts = re.split(r"\n## ", body)
    tresc = parts[0].strip()
    zalaczniki = ""
    otrzymuja = ""
    for part in parts[1:]:
        heading, _, rest = part.partition("\n")
        heading_lower = heading.strip().lower()
        rest = rest.strip()
        if heading_lower.startswith("załączniki"):
            zalaczniki = rest
        elif heading_lower.startswith("otrzymują"):
            otrzymuja = rest
    return {"tresc": tresc, "zalaczniki": zalaczniki, "otrzymuja": otrzymuja}


# ---------------------------------------------------------------------------
# Markdown → plain paragraphs
# ---------------------------------------------------------------------------

def md_to_paragraphs(md: str) -> list[tuple[str, str]]:
    """Convert a markdown snippet to a list of (style, text) tuples.

    The style hint is one of:
        "heading3"  — for `### ` or `#### ` lines (maps to style "Heading 3" /
                       falls back to bold normal if style is missing in DOCX).
        "quote"     — for blockquote lines (Obsidian callouts included).
        "list"      — for numbered / bulleted list items (flattened to body text
                       with bullet prefix).
        "normal"    — everything else (including emphasized lines).
    Blank lines are dropped; adjacent same-style paragraphs remain separate so
    that DOCX paragraph breaks match markdown intent.
    """
    paragraphs: list[tuple[str, str]] = []
    buffer: list[str] = []
    current_style: str = "normal"

    def flush():
        nonlocal buffer, current_style
        if buffer:
            text = " ".join(buffer).strip()
            if text:
                paragraphs.append((current_style, text))
        buffer = []

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            flush()
            continue
        stripped = line.strip()
        if stripped.startswith("#### "):
            flush()
            current_style = "heading3"
            buffer.append(_strip_md_inline(stripped[5:].strip()))
            flush()
            current_style = "normal"
            continue
        if stripped.startswith("### "):
            flush()
            current_style = "heading3"
            buffer.append(_strip_md_inline(stripped[4:].strip()))
            flush()
            current_style = "normal"
            continue
        if stripped.startswith("## "):
            # Should not happen inside tresc, but guard against it.
            flush()
            continue
        if stripped.startswith("> "):
            flush()
            current_style = "quote"
            buffer.append(_strip_md_inline(_strip_callout_prefix(stripped[2:])))
            flush()
            current_style = "normal"
            continue
        if stripped.startswith(">"):
            flush()
            current_style = "quote"
            buffer.append(_strip_md_inline(_strip_callout_prefix(stripped[1:].lstrip())))
            flush()
            current_style = "normal"
            continue
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            flush()
            current_style = "list"
            buffer.append(f"{m.group(1)}. {_strip_md_inline(m.group(2))}")
            flush()
            current_style = "normal"
            continue
        if stripped.startswith(("- ", "* ")):
            flush()
            current_style = "list"
            buffer.append(f"• {_strip_md_inline(stripped[2:])}")
            flush()
            current_style = "normal"
            continue
        if stripped.startswith("---"):
            flush()
            continue
        if stripped.startswith("<!--"):
            # Skip HTML comments (editorial notes in templates).
            continue
        # Default: body text — may be continuation.
        buffer.append(_strip_md_inline(stripped))

    flush()
    return paragraphs


def _strip_callout_prefix(text: str) -> str:
    # Obsidian callouts: `[!info]`, `[!warning] Title`, etc.
    match = re.match(r"^\[![a-zA-Z]+\]\s*(.*)$", text)
    if match:
        rest = match.group(1).strip()
        return rest or "(uwaga)"
    return text


def _strip_md_inline(text: str) -> str:
    # Remove **bold**, __bold__, *italic*, `code`, [text](url), highlights `==`.
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    text = re.sub(r"==([^=]+)==", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[\[([^\]|]+)(\|[^\]]+)?\]\]", r"\1", text)
    return text.strip()


# ---------------------------------------------------------------------------
# DOCX manipulation
# ---------------------------------------------------------------------------

def format_polish_date(value) -> str:
    if isinstance(value, dt.date):
        return f"{value.day} {POLISH_MONTHS[value.month]} {value.year}"
    if isinstance(value, str):
        # Accept ISO (YYYY-MM-DD) or already-formatted Polish date.
        iso_match = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", value.strip())
        if iso_match:
            y, m, d = map(int, iso_match.groups())
            return f"{d} {POLISH_MONTHS[m]} {y}"
        return value
    return str(value)


def collect_bookmarks(doc) -> dict[str, list]:
    """Return {bookmark_name: [bookmarkStart_element, ...]}."""
    bookmarks: dict[str, list] = {}
    for bm_start in doc.element.iter(qn("w:bookmarkStart")):
        name = bm_start.get(qn("w:name"))
        if name:
            bookmarks.setdefault(name, []).append(bm_start)
    return bookmarks


def replace_bookmark_text(bm_start, new_text: str) -> None:
    """Replace text covered by a bookmark with `new_text`.

    Finds the matching bookmarkEnd (by w:id) and replaces all w:t descendants
    between them with a single new w:t. Removes now-empty runs.
    """
    bm_id = bm_start.get(qn("w:id"))
    parent = bm_start.getparent()
    # Walk siblings until we find w:bookmarkEnd with matching id.
    # Bookmarks can span multiple siblings (runs, paragraphs); handle flat case
    # first by walking paragraph-level siblings via document order.
    doc_elem = bm_start
    while doc_elem.getparent() is not None:
        doc_elem = doc_elem.getparent()
    # Collect nodes between start and end.
    in_range = False
    nodes_in_range: list = []
    end_elem = None
    for node in doc_elem.iter():
        if node is bm_start:
            in_range = True
            continue
        if in_range:
            if node.tag == qn("w:bookmarkEnd") and node.get(qn("w:id")) == bm_id:
                end_elem = node
                break
            nodes_in_range.append(node)
    if end_elem is None:
        return
    # Strip text from all w:t nodes in range.
    text_nodes = [n for n in nodes_in_range if n.tag == qn("w:t")]
    if not text_nodes:
        # Bookmark has no text — insert a new run after bookmarkStart.
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        bm_start.addnext(run)
        return
    # Put new text into the first w:t, clear the rest.
    text_nodes[0].text = new_text
    text_nodes[0].set(qn("xml:space"), "preserve")
    for n in text_nodes[1:]:
        n.text = ""


def iter_all_paragraphs(doc) -> Iterable:
    """Yield every paragraph in body, headers, and footers."""
    yield from doc.paragraphs
    for section in doc.sections:
        for part_name in (
            "header", "first_page_header", "even_page_header",
            "footer", "first_page_footer", "even_page_footer",
        ):
            part = getattr(section, part_name, None)
            if part is None:
                continue
            yield from part.paragraphs


def paragraph_text_equal(paragraph, *prefixes: str) -> bool:
    text = paragraph.text.strip()
    return any(text.startswith(prefix) for prefix in prefixes)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace the paragraph's text while preserving paragraph properties.

    Removes all existing w:r children, then appends a single new run with the
    first run's rPr (if any) and the new text.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    # Capture original rPr from the first run, if any.
    first_rPr = None
    for r in paragraph._p.findall(qn("w:r")):
        if first_rPr is None:
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                first_rPr = deepcopy(rpr)
        paragraph._p.remove(r)
    # Remove any lingering bookmarkStart/End inside this paragraph — they are
    # re-emitted later by bookmark logic if needed. We keep them because they
    # may be necessary for downstream systems (EZD integrates via them).
    new_r = OxmlElement("w:r")
    if first_rPr is not None:
        new_r.append(first_rPr)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    new_r.append(new_t)
    paragraph._p.append(new_r)


def replace_all_text_inline(paragraph, mapping: dict[str, str]) -> None:
    """Replace exact string substrings in a paragraph's combined text.

    This rebuilds the paragraph's runs to hold a single run with the merged
    (replaced) text. Formatting from the first run is preserved.
    """
    original_text = paragraph.text
    if not any(key in original_text for key in mapping):
        return
    new_text = original_text
    for key, value in mapping.items():
        new_text = new_text.replace(key, value)
    if new_text == original_text:
        return
    replace_paragraph_text(paragraph, new_text)


def replace_all_text_ordered(paragraph, pairs: list[tuple[str, str]]) -> None:
    """Like replace_all_text_inline but honours the provided (key, value) order.

    Use for cases where one placeholder is a substring of another — the caller
    must put the longer form first so the composite replacement happens before
    its constituents are substituted.
    """
    original_text = paragraph.text
    if not any(key in original_text for key, _ in pairs):
        return
    new_text = original_text
    for key, value in pairs:
        new_text = new_text.replace(key, value)
    if new_text == original_text:
        return
    replace_paragraph_text(paragraph, new_text)


def _bookmark_range_text(bm_start, bm_id: str, doc) -> str:
    """Return the concatenated text inside a bookmark range (document-wide)."""
    doc_elem = bm_start
    while doc_elem.getparent() is not None:
        doc_elem = doc_elem.getparent()
    in_range = False
    parts: list[str] = []
    for node in doc_elem.iter():
        if node is bm_start:
            in_range = True
            continue
        if in_range:
            if node.tag == qn("w:bookmarkEnd") and node.get(qn("w:id")) == bm_id:
                break
            if node.tag == qn("w:t") and node.text:
                parts.append(node.text)
    return "".join(parts)


def insert_paragraphs_before(anchor_paragraph, entries: list[tuple[str, str]]) -> None:
    """Insert new paragraphs before `anchor_paragraph`, copying its base style.

    `entries` is a list of (style, text) tuples (see md_to_paragraphs).
    """
    base_pPr = anchor_paragraph._p.find(qn("w:pPr"))
    for style, text in entries:
        new_p = OxmlElement("w:p")
        if base_pPr is not None:
            new_p.append(deepcopy(base_pPr))
        run = OxmlElement("w:r")
        rpr = _style_run_properties(style)
        if rpr is not None:
            run.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        new_p.append(run)
        anchor_paragraph._p.addprevious(new_p)


def _style_run_properties(style: str):
    if style == "heading3":
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rpr.append(b)
        return rpr
    if style == "quote":
        rpr = OxmlElement("w:rPr")
        i = OxmlElement("w:i")
        rpr.append(i)
        return rpr
    return None


def remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


# ---------------------------------------------------------------------------
# Render pipeline
# ---------------------------------------------------------------------------

def render(
    *,
    template_path: Path,
    input_md_path: Path,
    output_path: Path,
    platforma_url: str | None = None,
) -> None:
    fm, sections = parse_md(input_md_path)

    _validate_frontmatter(fm)
    _validate_body(sections)

    shutil.copyfile(template_path, output_path)
    doc = Document(output_path)

    wykonawca = fm["wykonawca"]
    adres = fm.get("adres_wykonawcy") or "<<adres wykonawcy — uzupełnić przed wysłaniem>>"
    # Allow the letter frontmatter to use either `sygnatura_postepowania`
    # (drafting-pzp-letters convention) or `sygnatura` (analyzing-pzp-offers
    # convention) so users can paste analysis frontmatter without renaming.
    sygnatura_postepowania = fm.get("sygnatura_postepowania") or fm.get("sygnatura")
    sygnatura_pisma = fm.get(
        "sygnatura_pisma",
        f"{sygnatura_postepowania}.{fm['kod_pisma']}",
    )
    data_pisma_fmt = format_polish_date(fm["data_pisma"])
    stanowisko = fm.get("signatory_stanowisko", "")
    tytul = fm.get("signatory_tytul", "")
    imie_nazwisko = fm.get("signatory_imie_nazwisko", "")
    url_platformy = platforma_url or fm.get("platforma_url") or "<<URL platformy — uzupełnić przed wysłaniem>>"

    # Substitute body-level placeholders into the markdown string BEFORE it is
    # parsed into DOCX paragraphs. Inline replacement via `iter_all_paragraphs`
    # below runs against the template DOCX, not against paragraphs we inject
    # later, so `<<URL platformy>>` would otherwise survive in the body.
    for section_name, section_text in list(sections.items()):
        sections[section_name] = section_text.replace("<<URL platformy>>", url_platformy)

    # 1) Inline placeholders (primary mechanism) ---------------------------
    # The template uses `$...` placeholders inside EZD bookmark ranges. The
    # bookmarks are anchors for the EZD runtime to read/write metadata; for our
    # render pass the cheapest and safest approach is to substitute the text
    # directly. EZD bookmarks wrap the paragraph contents so the replaced text
    # remains inside the bookmark range. The longest composite key MUST be
    # applied before its constituent shorter keys (otherwise "$tytuł" would
    # match inside "$tytuł $imię i nazwisko" and leave the "$imię..." half
    # untouched).
    imie_tytul_value = f"{tytul} {imie_nazwisko}".strip()
    inline_map_ordered = [
        ("$sygnatura pisma", sygnatura_pisma),
        ("$DataPodpisu", data_pisma_fmt),
        ("$stanowisko", stanowisko),
        ("$tytuł $imię i nazwisko", imie_tytul_value),
        ("$imię i nazwisko", imie_nazwisko),
        ("$tytuł", tytul),
        # Body placeholder used across all letter templates. Substituted
        # regardless of whether the agent filled it into the .md body or not,
        # so a missing `--platforma` CLI argument surfaces the placeholder
        # text rather than silently leaving `<<URL platformy>>` in the DOCX.
        ("<<URL platformy>>", url_platformy),
    ]
    for para in list(iter_all_paragraphs(doc)):
        replace_all_text_ordered(para, inline_map_ordered)

    # 2) Bookmarks — only the empty ezdPracownikAtrybut2 needs a new run ---
    # All other EZD bookmarks already wrap text that was updated in step 1, so
    # the bookmark range content is correct. ezdPracownikAtrybut2 in the
    # template is a bookmark around an empty paragraph (serves as an EZD anchor
    # for a secondary attribute). We populate it with the signatory's degree
    # (tytuł) so that the rendered DOCX has a complete signature block even
    # when no EZD runtime is involved.
    bookmarks = collect_bookmarks(doc)
    for bm in bookmarks.get("ezdPracownikAtrybut2", []):
        # Only fill when the bookmark range contains no real text.
        bm_id = bm.get(qn("w:id"))
        existing_text = _bookmark_range_text(bm, bm_id, doc)
        if not existing_text.strip():
            replace_bookmark_text(bm, tytul)

    # 3) Adresat paragraph --------------------------------------------------
    addressee_text = f"{wykonawca}\n{adres}"
    # In DOCX we do not embed literal newlines in a single w:t — split into two
    # paragraphs.
    for para in list(doc.paragraphs):
        if "[adresat]" in para.text:
            lines = [line for line in addressee_text.split("\n") if line.strip()]
            if not lines:
                continue
            replace_paragraph_text(para, lines[0])
            # Chain extra paragraphs after the previously-inserted one so that
            # lines keep their original order (addnext inserts immediately
            # after the anchor; using `para` as anchor each iteration would
            # reverse the sequence for 3+ lines).
            anchor_p = para._p
            pPr = para._p.find(qn("w:pPr"))
            for extra in lines[1:]:
                new_p = OxmlElement("w:p")
                if pPr is not None:
                    new_p.append(deepcopy(pPr))
                run = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = extra
                run.append(t)
                new_p.append(run)
                anchor_p.addnext(new_p)
                anchor_p = new_p
            break

    # 4) Body text (Wstęp + Rozwinięcie placeholders → markdown content) ---
    body_paragraphs = md_to_paragraphs(sections["tresc"])
    anchor_body_paragraphs = [
        p for p in list(doc.paragraphs)
        if paragraph_text_equal(p, *BODY_PLACEHOLDER_PREFIXES)
    ]
    if anchor_body_paragraphs:
        first_anchor = anchor_body_paragraphs[0]
        insert_paragraphs_before(first_anchor, body_paragraphs)
        for anchor in anchor_body_paragraphs:
            remove_paragraph(anchor)

    # 5) Załączniki ---------------------------------------------------------
    zal_text = sections.get("zalaczniki", "").strip()
    zal_entries = md_to_paragraphs(zal_text) if zal_text else [("normal", "nie dotyczy")]
    for para in list(doc.paragraphs):
        if paragraph_text_equal(para, *ATTACHMENT_PLACEHOLDER_PREFIXES):
            insert_paragraphs_before(para, zal_entries)
            remove_paragraph(para)
            break

    # Replace trailing "……" in Załączniki section if present.
    for para in list(doc.paragraphs):
        if para.text.strip() in {"……", "…….", "…"}:
            remove_paragraph(para)
            break

    # 6) Otrzymują ----------------------------------------------------------
    otrz_text = sections.get("otrzymuja", "").strip()
    otrz_entries = md_to_paragraphs(otrz_text) if otrz_text else [("normal", f"{wykonawca}, {adres}"), ("normal", "a/a")]
    for para in list(doc.paragraphs):
        if paragraph_text_equal(para, *RECIPIENT_PLACEHOLDER_PREFIXES):
            insert_paragraphs_before(para, otrz_entries)
            remove_paragraph(para)
            break
    for para in list(doc.paragraphs):
        if para.text.strip() == "……":
            remove_paragraph(para)
            break

    doc.save(output_path)
    _warn_signatory(fm)


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

REQUIRED_FRONTMATTER_FIELDS = (
    "postepowanie",
    "zamawiajacy",
    "wykonawca",
    "typ_pisma",
    "kod_pisma",
    "data_pisma",
    "miejscowosc",
    "signatory_stanowisko",
    "signatory_tytul",
    "signatory_imie_nazwisko",
)

# Body placeholders that are substituted by the render pipeline itself and
# therefore should NOT trigger the body-placeholder validator.
_RENDERER_PROVIDED_PLACEHOLDERS = frozenset({"<<URL platformy>>"})

_PLACEHOLDER_PATTERN = re.compile(r"<<[^>]*>>")


def _validate_frontmatter(fm: dict) -> None:
    missing = [key for key in REQUIRED_FRONTMATTER_FIELDS if not fm.get(key)]
    if not (fm.get("sygnatura_postepowania") or fm.get("sygnatura")):
        missing.append("sygnatura_postepowania (or `sygnatura`)")
    if missing:
        raise ValueError(
            "Missing mandatory frontmatter keys: " + ", ".join(missing)
        )
    for key, value in fm.items():
        if isinstance(value, str) and _PLACEHOLDER_PATTERN.search(value):
            raise ValueError(
                f"Frontmatter field '{key}' still contains placeholder: {value!r}. "
                "Replace all <<...>> with real values before rendering."
            )


def _validate_body(sections: dict[str, str]) -> None:
    """Raise when the letter body still contains unfilled `<<...>>` tokens.

    Skips placeholders that the render pipeline substitutes itself (e.g.
    `<<URL platformy>>` — injected from the `--platforma` CLI argument or a
    fallback warning string).
    """
    offenders: list[tuple[str, str]] = []
    for section_name, section_text in sections.items():
        for match in _PLACEHOLDER_PATTERN.finditer(section_text):
            token = match.group(0)
            if token in _RENDERER_PROVIDED_PLACEHOLDERS:
                continue
            offenders.append((section_name, token))
    if offenders:
        first_section, first_token = offenders[0]
        count = len(offenders)
        raise ValueError(
            f"Letter body still contains {count} unfilled placeholder(s) — "
            f"first offender in section '{first_section}': {first_token!r}. "
            "Replace all <<...>> with real values before rendering."
        )


def _warn_signatory(fm: dict) -> None:
    source = fm.get("signatory_zrodlo", "").lower()
    if source == "memory":
        sys.stderr.write(
            "\n[UWAGA] Signatory pobrany z memory (signatory_zrodlo: memory). "
            "Zweryfikuj dane osoby podpisującej przed wysyłką pisma.\n"
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument(
        "--platforma",
        default=None,
        help="Optional URL of the procurement platform (substituted for "
             "'<<URL platformy>>' placeholders in the body text).",
    )
    args = parser.parse_args(argv)

    if not args.template.is_file():
        parser.error(f"Template file not found: {args.template}")
    if not args.input.is_file():
        parser.error(f"Input markdown not found: {args.input}")
    args.output.parent.mkdir(parents=True, exist_ok=True)

    render(
        template_path=args.template,
        input_md_path=args.input,
        output_path=args.output,
        platforma_url=args.platforma,
    )
    sys.stdout.write(f"Wygenerowano: {args.output}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
